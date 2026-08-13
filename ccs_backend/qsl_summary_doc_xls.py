import os
from datetime import datetime, timezone

import config
import handle_db
from handle_db import Log


def formatTimestamp(timestamp):
    """A napló UTC időbélyegéből olvasható dátum + idő: '2026-08-01 14:23'."""
    if not timestamp:
        return "-"
    return datetime.fromtimestamp(int(timestamp), timezone.utc).strftime("%Y-%m-%d %H:%M")


# A magyar amatőr hívójelek prefixei.
HUNGARIAN_PREFIXES = ("HA", "HG")


def getQsosByParticipant(minValidQso=1, onlyHungarian=False):
    """Résztvevőnként a QSO-i: sáv, mód, operátor (időrendben).

    Egyetlen lekérdezés az egész naplóra, a csoportosítás pedig a normalizált
    hívójel szerint megy – ugyanaz a szabály, mint a handle_db.query()-ben
    (HA1DC és DL_HA1DC egy állomás), de résztvevőnkénti külön lekérdezés
    (1356 db) nélkül.

    minValidQso: csak azok a résztvevők maradnak benne, akiknek legalább ennyi
    KÜLÖNBÖZŐ (sáv, mód) párosításban van QSO-juk – ugyanaz a számolás, mint a
    diplománál (minValidQso=3 -> a diplomát szerzettek). Ugyanazon a sávon és
    módon a további QSO-k nem számítanak bele. Az alapértelmezett 1 mindenkit
    benne hagy.

    onlyHungarian: csak a magyar (HA / HG kezdetű) hívójelek. A szűrés a
    normalizált alakon fut, tehát a DL_HA1DC-féle bejegyzés is magyarnak
    számít, mert az igazi hívójele HA1DC.

    Visszatérés: {hívójel: [{"band":..., "mode":..., "operator":...}, ...]}
    """
    with handle_db._db() as s:
        rows = (s.query(Log.log_timestamp_utc,Log.callsign, Log.band, Log.mode, Log.rst_sent, Log.local_operator)
                 .order_by(Log.log_timestamp_utc)
                 .all())

    qsos = dict()
    seenBandModes = dict()
    for timestamp, callsign, band, mode, rst_sent, operator in rows:
        core = handle_db._normalize_callsign(callsign)
        if not core:
            continue

        if onlyHungarian and not core.startswith(HUNGARIAN_PREFIXES):
            continue

        # A sorok időrendben jönnek, így az adott (sáv, mód) ELSŐ QSO-ja az
        # érvényes, a későbbiek már csak ismétlések – a diplomába nem
        # számítanak bele.
        seen = seenBandModes.setdefault(core, set())
        repeated = (band, mode) in seen
        seen.add((band, mode))

        qsos.setdefault(core, []).append({"datetime": int(timestamp),
                                          "band": band,
                                          "mode": mode,
                                          "rst_sent": rst_sent,
                                          "operator": operator,
                                          "repeated": repeated})

    # A szűrés az EGYEDI (sáv, mód) párokat nézi, nem a QSO-k darabszámát:
    # 5 QSO ugyanazon a sávon és módon 1 érvényesnek számít.
    return {callsign: qsoList for callsign, qsoList in sorted(qsos.items())
            if len({(i["band"], i["mode"]) for i in qsoList}) >= minValidQso}


# A QSO-táblázat fejléce – a konzolra írt oszlopokkal megegyezően. Az ismétlés
# jelölése formátumonként más: az XLS külön oszlopban írja ki, a DOC áthúzza az
# érintett sort.
QSO_HEADER = ("Dátum (UTC)", "Sáv", "Mód", "RST küldött", "Operátor")
REPEAT_HEADER = "Ismétlés"


def _qsoRow(qso):
    """Egy QSO a táblázatok soraként, a konzolos kiírással azonos tartalommal."""
    return (formatTimestamp(qso["datetime"]),
            qso["band"] or "-",
            qso["mode"] or "-",
            qso.get("rst_sent") or "-",
            (qso.get("operator") or "-").upper())


def _validQsoNr(qsos):
    """Az érvényes (nem ismétlés) QSO-k száma – ez számít a diplomába."""
    return sum(1 for qso in qsos if not qso["repeated"])


def generateDoc(qsosByParticipant, outPath, title=None):
    """Word (.docx) dokumentum a résztvevők QSO-iból.

    Résztvevőnként egy cím (hívójel, QSO-szám, érvényes QSO-szám) és alatta a
    QSO-k táblázata – ugyanaz a bontás, mint a konzolos listában.

    Az ismétlés (már meglévő sáv-mód páros) itt nem külön oszlop: az érintett
    sor át van húzva.

    A cím alapértelmezésben az aktiválás hívójelével készül (konfigurációból).
    """
    from docx import Document
    from docx.shared import Pt

    if title is None:
        title = f"{config.getActivationCallsign().upper()} – QSO összesítő"

    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph(f"Készült: {formatTimestamp(handle_db.getCurrentUtcTs())} UTC – "
                           f"{len(qsosByParticipant)} résztvevő, "
                           f"{sum(len(i) for i in qsosByParticipant.values())} QSO")

    for callsign, qsos in qsosByParticipant.items():
        document.add_heading(f"{callsign} ({len(qsos)} qso, {_validQsoNr(qsos)} érvényes)", level=1)

        table = document.add_table(rows=1, cols=len(QSO_HEADER))
        table.style = "Table Grid"
        for cell, header in zip(table.rows[0].cells, QSO_HEADER):
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        for qso in qsos:
            for cell, value in zip(table.add_row().cells, _qsoRow(qso)):
                cell.text = value
                font = cell.paragraphs[0].runs[0].font
                font.size = Pt(9)
                # Az ismétlődő sáv-mód páros nem számít a diplomába: áthúzva.
                font.strike = qso["repeated"]

    os.makedirs(os.path.dirname(os.path.abspath(outPath)), exist_ok=True)
    document.save(outPath)
    return outPath


def generateXls(qsosByParticipant, outPath):
    """Excel (.xlsx) munkafüzet a résztvevők QSO-iból.

    Két munkalap: a 'QSO-k' minden QSO-t soronként (első oszlopban a hívójel,
    hogy szűrhető/rendezhető legyen), az 'Összesítő' pedig résztvevőnként a
    QSO- és az érvényes QSO-számot.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Border, Font, Side
    from openpyxl.utils import get_column_letter

    workbook = Workbook()
    boldFont = Font(bold=True)
    thinSide = Side(style="thin")

    qsoSheet = workbook.active
    qsoSheet.title = "QSO-k"
    qsoSheet.append(("Hívójel",) + QSO_HEADER + (REPEAT_HEADER,))

    summarySheet = workbook.create_sheet("Összesítő")
    summarySheet.append(("Hívójel", "QSO", "Érvényes QSO"))

    callsignBlocks = list()
    for callsign, qsos in qsosByParticipant.items():
        startRow = qsoSheet.max_row + 1
        for qso in qsos:
            qsoSheet.append((callsign,) + _qsoRow(qso)
                            + ("ismétlés" if qso["repeated"] else "",))
        callsignBlocks.append((startRow, qsoSheet.max_row))

        summarySheet.append((callsign, len(qsos), _validQsoNr(qsos)))

    # Egy hívójel sorai köré keret, hogy a blokkok jól elváljanak egymástól.
    lastColumn = qsoSheet.max_column
    for startRow, endRow in callsignBlocks:
        for row in qsoSheet.iter_rows(min_row=startRow, max_row=endRow,
                                      min_col=1, max_col=lastColumn):
            for cell in row:
                cell.border = Border(
                    top=thinSide if cell.row == startRow else None,
                    bottom=thinSide if cell.row == endRow else None,
                    left=thinSide if cell.column == 1 else None,
                    right=thinSide if cell.column == lastColumn else None)

    for sheet, widths in ((qsoSheet, (12, 18, 8, 8, 10, 10)),
                          (summarySheet, (12, 8, 14))):
        for cell in sheet[1]:
            cell.font = boldFont
        for idx, width in enumerate(widths, start=1):
            sheet.column_dimensions[get_column_letter(idx)].width = width
        # A fejléc maradjon látható görgetéskor, és lehessen rá szűrni.
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions

    os.makedirs(os.path.dirname(os.path.abspath(outPath)), exist_ok=True)
    workbook.save(outPath)
    return outPath


if __name__ == "__main__":
    qsosByParticipant = getQsosByParticipant(minValidQso=3, onlyHungarian=True)

    for callsign, qsos in qsosByParticipant.items():
        validNr = sum(1 for qso in qsos if not qso["repeated"])
        print(f"{callsign} ({len(qsos)} qso, {validNr} érvényes)")
        for qso in qsos:
            print(f"    {formatTimestamp(qso['datetime']):16s} {qso['band'] or '-':6s} {qso['mode'] or '-':6s} {qso['rst_sent'] or '-':6s} {qso['operator'].upper() or '-':8s}"
                  f"{'  ismétlés' if qso['repeated'] else ''}")

    #print()
    #print("doc:", generateDoc(qsosByParticipant, "./tmp/qso_summary.docx"))
    #print("xls:", generateXls(qsosByParticipant, "./tmp/qso_summary.xlsx"))


